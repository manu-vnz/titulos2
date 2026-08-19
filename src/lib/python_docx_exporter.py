import docx
import re
import copy
import os
import glob
from lxml import etree


# ─── Default preview positions (% of canvas) that match the hardcoded CSS defaults.
# Used as the reference point for computing deltas when the user drags fields.
DEFAULT_PREVIEW_POSITIONS = {
    'plantel':                 {'top': 4.0,   'left': 21.0},
    'codigo_plantel':          {'top': 1.5,   'left': 8.2},
    'titulo_otorgado':         {'top': 3.9,   'left': 9.2},
    'plan_estudio':            {'top': 2.3,   'left': 24.6},
    'nombre_estudiante':       {'top': 8.8,   'left': 15.9},
    'cedula_estudiante':       {'top': 11.2,  'left': 22.6},
    'lugar_nacimiento':        {'top': 1.2,   'left': 13.8},
    'fecha_nacimiento':        {'top': 16.0,  'left': 9.1},
    'lugar_fecha_expedicion':  {'top': 20.7,  'left': 24.2},
    'ano_egreso':              {'top': 22.9,  'left': 13.8},
    'coordinador_nombre':      {'top': 18.5,  'left': 6.8},
    'coordinador_cedula':      {'top': 22.9,  'left': 31.3},
    'funcionario_nombre':      {'top': 20.6,  'left': 32.4},
    'funcionario_cedula':      {'top': 20.7,  'left': 4.9},
    'director_nombre':         {'top': 20.6,  'left': 66.7},
    'director_cedula':         {'top': 22.9,  'left': 62.4},
}

# ─── Mapping from field key to the original text in the gold template.
# Used to identify which VML shape corresponds to which field.
FIELD_TO_TEMPLATE_TEXT = {
    'plantel':                 'COMPLEJO EDUCATIVO',
    'nombre_estudiante':       'JESUS MANUEL VARGAS NOGUERA',
    'cedula_estudiante':       'V 33.479.449',
    'lugar_fecha_expedicion':  'CARABOBO, VALENCIA, 17 DE JULIO',
    'fecha_nacimiento':        '09 DE JULIO DE 2009',
    'ano_egreso':              '2026',
    'titulo_otorgado':         'BACHILLER',
    'codigo_plantel':          'S0163D0814',
    'plan_estudio':            'MEDIA GENERAL',
    'lugar_nacimiento':        'VENEZUELA, CARABOBO, MUNICIPIO',
    'director_cedula':         'V 18.361.899',
    'coordinador_cedula':      'V 13.601.460',
    'funcionario_cedula':      'V 9.445.225',
    'director_nombre':         'JOHN DANIEL ZAPATA',
    'coordinador_nombre':      'ALBERTO',
    'funcionario_nombre':      'WILMER',
}

# Page dimensions in points (Letter landscape 11" x 8.5")
PAGE_WIDTH_PT = 792.0
PAGE_HEIGHT_PT = 612.0

# Scale factors: how many pt per 1% of the canvas
SCALE_X = PAGE_WIDTH_PT / 100.0   # 7.92 pt per %
SCALE_Y = PAGE_HEIGHT_PT / 100.0  # 6.12 pt per %


def find_gold_template(plantillas_dir):
    """
    Busca la plantilla dorada oficial 'JESUS MANUEL VARGAS NOGUERA COMPLEJO EDUCATIVO RUIZPINEDA I 2026.docx'.
    Ignora cualquier archivo que contenga 'chueco' o que sea una exportación previa.
    """
    gold_official = os.path.join(plantillas_dir, "JESUS MANUEL VARGAS NOGUERA COMPLEJO EDUCATIVO RUIZPINEDA I 2026.docx")
    if os.path.exists(gold_official):
        return gold_official

    docx_files = glob.glob(os.path.join(plantillas_dir, '*.docx'))
    valid_files = [
        f for f in docx_files
        if not re.search(r'chueco|consolidado|impresion|output|runner', f, re.IGNORECASE)
    ]
    if valid_files:
        valid_files.sort(key=os.path.getmtime, reverse=True)
        return valid_files[0]

    if docx_files:
        return docx_files[0]

    raise FileNotFoundError(f"No se encontró ninguna plantilla base .docx en: {plantillas_dir}")


import unicodedata

def strip_accents(s):
    return ''.join(c for c in unicodedata.normalize('NFD', s) if unicodedata.category(c) != 'Mn')

def reemplazar_texto_en_parrafo(paragraph, buscar, reemplazar):
    """
    Unifica todos los runs de un párrafo para preservar el formato, la posición y la alineación.
    Sustituye la cadena 'buscar' por 'reemplazar' sin fragmentar los nodos <w:t>.
    Aplica negrita (bold = True) a todo el texto reemplazado.
    Insensible a tildes y mayúsculas/minúsculas.
    """
    full_text = paragraph.text
    if not full_text or not buscar:
        return

    # Try exact match first
    if buscar in full_text:
        new_text = full_text.replace(buscar, str(reemplazar or ''))
        if paragraph.runs:
            paragraph.runs[0].text = new_text
            paragraph.runs[0].bold = True
            for i in range(1, len(paragraph.runs)):
                paragraph.runs[i].text = ""
        return

    # Fallback to accent-insensitive match
    norm_full = strip_accents(full_text.upper())
    norm_buscar = strip_accents(buscar.upper())

    if norm_buscar in norm_full:
        start_idx = norm_full.find(norm_buscar)
        end_idx = start_idx + len(norm_buscar)
        new_text = full_text[:start_idx] + str(reemplazar or '') + full_text[end_idx:]
        if paragraph.runs:
            paragraph.runs[0].text = new_text
            paragraph.runs[0].bold = True
            for i in range(1, len(paragraph.runs)):
                paragraph.runs[i].text = ""

def resolve_field_pos(matched_field, field_positions):
    if not field_positions:
        return None
    aliases = [
        matched_field,
        'estudiante_nombre' if matched_field == 'nombre_estudiante' else matched_field,
        'estudiante_cedula' if matched_field == 'cedula_estudiante' else matched_field,
        'año_egreso' if matched_field == 'ano_egreso' else matched_field,
        'firmante_' + matched_field,
        matched_field.replace('firmante_', '') if matched_field.startswith('firmante_') else matched_field,
    ]
    for a in aliases:
        if a in field_positions:
            return field_positions[a]
    return None

def aplicar_posiciones_vml(doc, field_positions):
    """
    Applies position deltas from the frontend canvas to VML shapes in the DOCX.
    """
    if not field_positions:
        return

    shapes = doc.element.xpath('//*[local-name()="shape"]')
    
    for shape in shapes:
        style = shape.get('style', '')
        if not style:
            continue
        
        txbx_ps = shape.xpath('.//*[local-name()="txbxContent"]//*[local-name()="p"]')
        shape_text = ''
        for tp in txbx_ps:
            p = docx.text.paragraph.Paragraph(tp, doc)
            if p.text and p.text.strip():
                shape_text = p.text.strip()
                break
        
        if not shape_text:
            continue
        
        matched_field = None
        for field_key, search_text in FIELD_TO_TEMPLATE_TEXT.items():
            if search_text.upper() in shape_text.upper():
                matched_field = field_key
                break
        
        if not matched_field:
            continue
        
        new_pos = resolve_field_pos(matched_field, field_positions)
        if not new_pos:
            continue
        
        default_pos = DEFAULT_PREVIEW_POSITIONS.get(matched_field)
        if not default_pos:
            continue
        
        new_top = float(new_pos.get('top', default_pos['top']))
        new_left = float(new_pos.get('left', default_pos['left']))
        
        delta_top_pct = new_top - default_pos['top']
        delta_left_pct = new_left - default_pos['left']
        
        if abs(delta_top_pct) < 0.05 and abs(delta_left_pct) < 0.05:
            continue
        
        delta_top_pt = delta_top_pct * SCALE_Y
        delta_left_pt = delta_left_pct * SCALE_X
        
        ml_match = re.search(r'margin-left:([\d.-]+)pt', style)
        mt_match = re.search(r'margin-top:([\d.-]+)pt', style)
        
        if ml_match and mt_match:
            current_ml = float(ml_match.group(1))
            current_mt = float(mt_match.group(1))
            
            new_ml = max(0, current_ml + delta_left_pt)
            new_mt = max(0, current_mt + delta_top_pt)
            
            new_style = re.sub(
                r'margin-left:[\d.-]+pt',
                'margin-left:%.2fpt' % new_ml,
                style
            )
            new_style = re.sub(
                r'margin-top:[\d.-]+pt',
                'margin-top:%.2fpt' % new_mt,
                new_style
            )
            
            shape.set('style', new_style)


def exportar_diploma_docx(template_path, output_path, datos, field_positions=None):
    """
    Motor de exportación .docx con unificación de runs y reemplazo limpio in-place.
    Aplica negrita (bold = True) a todo el texto del documento.
    Opcionalmente aplica posiciones personalizadas de campos VML.
    """
    doc = docx.Document(template_path)

    # Apply custom VML positions BEFORE text replacement (so shapes can be found by template text)
    if field_positions:
        aplicar_posiciones_vml(doc, field_positions)

    # Build the replacements dict — use ONLY literal text replacements from the gold template.
    nombre_nuevo = datos.get('estudiante_nombre') or datos.get('nombre_estudiante') or datos.get('nombre') or ''
    cedula_nueva = datos.get('estudiante_cedula') or datos.get('cedula_estudiante') or datos.get('cedula') or ''
    plantel_nuevo = datos.get('plantel') or datos.get('zona_educativa_plantel') or ''
    codigo_nuevo = datos.get('codigo_plantel') or ''
    titulo_nuevo = datos.get('titulo_otorgado') or ''
    plan_nuevo = datos.get('plan_estudio') or ''
    lugar_nac_nuevo = datos.get('lugar_nacimiento') or ''
    fecha_nac_nueva = datos.get('fecha_nacimiento') or ''
    expedicion_nueva = datos.get('lugar_fecha_expedicion') or ''
    ano_egreso_nuevo = str(datos.get('año_egreso') or datos.get('ano_egreso') or '2026')
    
    director_nom = datos.get('firmante_director_nombre') or datos.get('director_nombre') or 'JOHN DANIEL ZAPATA MIRELES'
    director_ci = datos.get('firmante_director_cedula') or datos.get('director_cedula') or 'V 18.361.899'
    coord_nom = datos.get('firmante_coordinador_nombre') or datos.get('coordinador_nombre') or 'JOS\u00c9 ALBERTO RU\u00cdZ \u00c1LVAREZ'
    coord_ci = datos.get('firmante_coordinador_cedula') or datos.get('coordinador_cedula') or 'V 13.601.460'
    func_nom = datos.get('firmante_funcionario_nombre') or datos.get('funcionario_nombre') or 'WILMER JOS\u00c9 LUGO RODR\u00cdGUEZ'
    func_ci = datos.get('firmante_funcionario_cedula') or datos.get('funcionario_cedula') or 'V 9.445.225'

    # Ordered from LONGEST to SHORTEST to prevent substring corruption.
    reemplazos = []
    
    if expedicion_nueva:
        reemplazos.append(("CARABOBO, VALENCIA, 17 DE JULIO DE 2026", expedicion_nueva))
    if lugar_nac_nuevo:
        reemplazos.append(("VENEZUELA, CARABOBO, MUNICIPIO NAGUANAGUA", lugar_nac_nuevo))
    if plan_nuevo:
        reemplazos.append(("EDUCACI\u00d3N MEDIA GENERAL, 31059", plan_nuevo))
        reemplazos.append(("EDUCACION MEDIA GENERAL, 31059", plan_nuevo))
    if nombre_nuevo:
        reemplazos.append(("JESUS MANUEL VARGAS NOGUERA", nombre_nuevo))
    if plantel_nuevo:
        reemplazos.append(("COMPLEJO EDUCATIVO RU\u00cdZ PINEDA I", plantel_nuevo))
        reemplazos.append(("COMPLEJO EDUCATIVO RUIZ PINEDA I", plantel_nuevo))
    if coord_nom:
        reemplazos.append(("JOS\u00c9 ALBERTO RU\u00cdZ \u00c1LVAREZ", coord_nom))
        reemplazos.append(("JOSE ALBERTO RUIZ ALVAREZ", coord_nom))
        reemplazos.append(("JOS\u00c9 ALBERTO RUIZ \u00c1LVAREZ", coord_nom))
    if func_nom:
        reemplazos.append(("WILMER JOS\u00c9 LUGO RODR\u00cdGUEZ", func_nom))
        reemplazos.append(("WILMER JOSE LUGO RODRIGUEZ", func_nom))
        reemplazos.append(("WILMER JOS\u00c9 LUGO RODRIGUEZ", func_nom))
    if director_nom:
        reemplazos.append(("JOHN DANIEL ZAPATA MIRELES", director_nom))
    if fecha_nac_nueva:
        reemplazos.append(("09 DE JULIO DE 2009", fecha_nac_nueva))
    if cedula_nueva:
        reemplazos.append(("V 33.479.449", cedula_nueva))
    if director_ci:
        reemplazos.append(("V 18.361.899", director_ci))
    if coord_ci:
        reemplazos.append(("V 13.601.460", coord_ci))
    if func_ci:
        reemplazos.append(("V 9.445.225", func_ci))
    if codigo_nuevo:
        reemplazos.append(("S0163D0814", codigo_nuevo))
    if titulo_nuevo:
        reemplazos.append(("BACHILLER", titulo_nuevo))
    if ano_egreso_nuevo:
        reemplazos.append(("2026", ano_egreso_nuevo))

    # Collect ALL paragraphs: body + tables + VML textboxes
    all_paragraphs = []
    
    for p in doc.paragraphs:
        all_paragraphs.append(p)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    all_paragraphs.append(p)
    
    shape_p_elements = doc.element.xpath(
        './/*[local-name()="txbxContent"]//*[local-name()="p"]'
    )
    for sp in shape_p_elements:
        all_paragraphs.append(docx.text.paragraph.Paragraph(sp, doc))
    
    # Apply replacements in order
    for buscar, reemplazar_val in reemplazos:
        for p in all_paragraphs:
            reemplazar_texto_en_parrafo(p, buscar, reemplazar_val)

    # Forzar negrita en absolutamente todos los párrafos y sus runs
    for p in all_paragraphs:
        for run in p.runs:
            if run.text and run.text.strip():
                run.bold = True

    doc.save(output_path)
    return doc


def generar_consolidado_multi_estudiante(template_path, students_list, output_path, field_positions=None):
    """
    Genera documento consolidado multi-estudiante.
    
    Strategy: Generate each student's page as a separate full document from the 
    gold template, then merge them using proper section breaks so each page 
    preserves the exact same layout, VML positioning, and formatting.
    """
    if not students_list:
        raise ValueError("La lista de estudiantes está vacía.")

    # Generate student 1 as the master document
    exportar_diploma_docx(template_path, output_path, students_list[0], field_positions)

    if len(students_list) <= 1:
        print("Documento consolidado generado exitosamente: %s" % output_path)
        return

    # For multiple students: generate each one separately, then merge via XML
    master_doc = docx.Document(output_path)
    
    # Get the section properties from the master doc for cloning
    master_body = master_doc.element.body
    master_sect_pr = master_body.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr')
    
    for s_idx in range(1, len(students_list)):
        student = students_list[s_idx]
        
        # Generate this student's page from the ORIGINAL template
        tmp_sub_path = output_path + "_sub_%d.docx" % s_idx
        exportar_diploma_docx(template_path, tmp_sub_path, student, field_positions)
        sub_doc = docx.Document(tmp_sub_path)
        
        # Add a section break (next page) to the master document.
        # We do this by adding a paragraph with sectPr that clones the master section properties.
        nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        
        # Create a section-break paragraph before appending the new page content
        break_para = etree.SubElement(master_body, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p')
        break_ppr = etree.SubElement(break_para, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr')
        
        # Clone the master section properties into this paragraph's pPr
        # This creates a "next page" section break
        if master_sect_pr is not None:
            cloned_sect = copy.deepcopy(master_sect_pr)
            # Ensure it's a "nextPage" type section break
            sect_type = cloned_sect.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type')
            if sect_type is None:
                sect_type = etree.SubElement(cloned_sect, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type')
            sect_type.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 'nextPage')
            break_ppr.append(cloned_sect)
        
        # Insert the break paragraph BEFORE the final sectPr
        if master_sect_pr is not None:
            master_sect_pr.addprevious(break_para)
        
        # Now append all body elements from the sub-document (except its final sectPr)
        sub_body = sub_doc.element.body
        for element in list(sub_body):
            tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag
            if tag == 'sectPr':
                continue  # skip the sub-document's final section properties
            cloned = copy.deepcopy(element)
            if master_sect_pr is not None:
                master_sect_pr.addprevious(cloned)
            else:
                master_body.append(cloned)
        
        # Cleanup temp file
        try:
            if os.path.exists(tmp_sub_path):
                os.remove(tmp_sub_path)
        except Exception:
            pass

    master_doc.save(output_path)
    print("Documento consolidado generado exitosamente: %s" % output_path)


if __name__ == "__main__":
    plantillas_dir = r"F:\intento 2\plantillas"
    template_path = find_gold_template(plantillas_dir)
    print("Plantilla dorada seleccionada: %s" % template_path)

    test_students = [
        {
            'plantel': 'COMPLEJO EDUCATIVO RU\u00cdZ PINEDA I',
            'codigo_plantel': 'S0163D0814',
            'titulo_otorgado': 'BACHILLER',
            'plan_estudio': 'EDUCACI\u00d3N MEDIA GENERAL, 31059',
            'estudiante_nombre': 'JESUS MANUEL VARGAS NOGUERA',
            'estudiante_cedula': 'V 33.479.449',
            'lugar_nacimiento': 'VENEZUELA, CARABOBO, MUNICIPIO NAGUANAGUA',
            'fecha_nacimiento': '09 DE JULIO DE 2009',
            'lugar_fecha_expedicion': 'CARABOBO, VALENCIA, 17 DE JULIO DE 2026',
            'a\u00f1o_egreso': '2026',
            'firmante_director_nombre': 'JOHN DANIEL ZAPATA MIRELES',
            'firmante_director_cedula': 'V 18.361.899',
            'firmante_coordinador_nombre': 'JOS\u00c9 ALBERTO RU\u00cdZ \u00c1LVAREZ',
            'firmante_coordinador_cedula': 'V 13.601.460',
            'firmante_funcionario_nombre': 'WILMER JOS\u00c9 LUGO RODR\u00cdGUEZ',
            'firmante_funcionario_cedula': 'V 9.445.225',
        },
    ]
    generar_consolidado_multi_estudiante(
        template_path, test_students,
        r"F:\intento 2\test_python_docx_consolidado.docx"
    )
